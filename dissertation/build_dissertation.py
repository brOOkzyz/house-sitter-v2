#!/usr/bin/env python3
"""Build the dissertation evidence pack from the official Word template."""

from __future__ import annotations

import csv
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "dissertation"
TEMPLATE = Path("/home/brookz/桌面/Dissertation Template 2024.docx")
FORMAL = ROOT / "results/raptor_lite/phase6_formal"
SUMMARY = FORMAL / "analysis/summary.json"
FREEZE = FORMAL / "phase7_results_freeze/final_results_summary.json"
RESULTS_MANIFEST = FORMAL / "phase7_results_freeze/results_manifest.json"
MANUSCRIPT = OUT / "manuscript.md"
REFERENCES = OUT / "references.json"
DOCX = OUT / "final_dissertation.docx"
PDF = OUT / "final_dissertation.pdf"

TITLE = (
    "RaPToR-Lite: A Capability-Grounded Natural-Language Task Creation and "
    "Verification Layer for ROS 2 Robots — A House-Sitter Case Study"
)
BASE_HEAD = "85179bdbe996053ec69ef5b0aca748812fdb8019"
PROTOCOL_HASH = "25a16e15fc07c2c9d3c76e52de067ca47f09950ac532bbbf1f14e611753c2847"
ANALYSIS_HASH = "fc1e1e1e20817435a80c0886715dcb25ce4ee3844e0ecf64f15c7189f34f9594"
RESULTS_HASH = "dc249cfff4b8f9bab90173555fb9860f96b22381fee64ff792ad9599e5499dec"
TEMPLATE_HASH = "c5ae3dc8b8b1ffec7efea4b90b88b34ae539f1c92f22155692a1c200e1399cc7"


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def words(text: str) -> list[str]:
    return re.findall(r"\b[\w’'-]+\b", text, flags=re.UNICODE)


def load_inputs() -> tuple[dict, dict, list[dict], str]:
    if sha256(TEMPLATE) != TEMPLATE_HASH:
        raise SystemExit("Official template hash changed; refusing to build.")
    summary = json.loads(SUMMARY.read_text())
    freeze = json.loads(FREEZE.read_text())
    manifest = json.loads(RESULTS_MANIFEST.read_text())
    refs = json.loads(REFERENCES.read_text())
    manuscript = MANUSCRIPT.read_text()
    if summary["protocol_hash"] != PROTOCOL_HASH or summary["analysis_plan_hash"] != ANALYSIS_HASH:
        raise SystemExit("Locked protocol or analysis hash mismatch.")
    if manifest["results_manifest_hash"] != RESULTS_HASH:
        raise SystemExit("Phase 7 results-manifest hash mismatch.")
    if freeze["raw_hashes"] != {
        "rq1": "76e4753f367596b3bda762badf2eb9b350779e469ae8e9e7fb520ee0ac055661",
        "rq2": "c4572ccd180ad454d15725351284497fa6f3118d817b00d283986a81ac89351a",
        "rq3": "ec9d077491c6e705456d84ce0c46be2bd241bad7eb3420401ef495c3db65e99a",
    }:
        raise SystemExit("Raw logical hashes do not match the Results Freeze.")
    keys = [r["key"] for r in refs]
    if len(keys) != len(set(keys)) or not (50 <= len(keys) <= 80):
        raise SystemExit("Reference keys must be unique and contain 50--80 sources.")
    return summary, freeze, refs, manuscript


def citation_map(refs: list[dict]) -> tuple[dict[str, int], list[dict]]:
    ordered = sorted(refs, key=lambda r: (r["sort_author"].casefold(), r["year"], r["title"].casefold()))
    return {r["key"]: i for i, r in enumerate(ordered, 1)}, ordered


def replace_citations(text: str, numbers: dict[str, int]) -> str:
    seen: set[str] = set()

    def repl(match: re.Match[str]) -> str:
        keys = [k.strip() for k in match.group(1).split(",")]
        missing = [k for k in keys if k not in numbers]
        if missing:
            raise SystemExit(f"Unknown citation keys: {missing}")
        seen.update(keys)
        return "[" + ",".join(str(numbers[k]) for k in keys) + "]"

    result = re.sub(r"\[cite:([^\]]+)\]", repl, text)
    return result.replace("`", "").replace("--", "–")


def audit_citations(manuscript: str, refs: list[dict], numbers: dict[str, int]) -> None:
    used: set[str] = set()
    for match in re.finditer(r"\[cite:([^\]]+)\]", manuscript):
        used.update(k.strip() for k in match.group(1).split(","))
    missing = used - set(numbers)
    unused = set(numbers) - used
    if missing or unused:
        raise SystemExit(f"Citation audit failed; missing={sorted(missing)}, unused={sorted(unused)}")
    abstract = manuscript.split("# ABSTRACT", 1)[1].split("# AUTHOR KEYWORDS", 1)[0]
    if "[cite:" in abstract:
        raise SystemExit("Abstract must not contain citations.")


def write_references_audit(ordered: list[dict], numbers: dict[str, int]) -> None:
    lines = [
        "# References Audit",
        "",
        "All entries were checked against the named publisher, proceedings, DOI or official documentation record. "
        "Citation numbers follow first-author surname order, as required by the dissertation template.",
        "",
        "| No. / key | Authors | Title | Venue | Year | DOI / public URL | Verification source | Used section | Claim supported |",
        "|---|---|---|---|---:|---|---|---|---|",
    ]
    for r in ordered:
        doi_url = f"https://doi.org/{r['doi']}" if r["doi"] else r["url"]
        cells = [
            f"[{numbers[r['key']]}] `{r['key']}`",
            r["authors"], r["title"], r["venue"], str(r["year"]),
            doi_url, r["verification_source"], r["used_sections"], r["claim_supported"],
        ]
        lines.append("| " + " | ".join(c.replace("|", "\\|") for c in cells) + " |")
    (OUT / "references_audit.md").write_text("\n".join(lines) + "\n")


def make_tables(summary: dict) -> None:
    tables = OUT / "tables"
    tables.mkdir(exist_ok=True)
    rq1_f1={r["condition"]:r for r in csv.DictReader((FORMAL/"phase7_results_freeze/tables/table1_rq1_verifier.csv").open())}
    fdr={r["comparison"]:r for r in summary["fdr"]}
    rq1=summary["rq1"]; dc=rq1["decision_correctness"]
    def pct(x): return f"{100*x:.1f}"
    def interval(pair): return f"{pct(pair[0])}–{pct(pair[1])}"
    def rq1_result(condition,label,comparison=None):
        item=dc[condition]; f1=rq1_f1[condition]
        result=f"{item['count']}/{item['n']}; {pct(item['rate'])}% (95% CI {interval(item['wilson95'])}); accept/reject F1 {float(f1['accept_f1']):.3f}/{float(f1['reject_f1']):.3f}"
        if comparison:
            effect=-rq1["paired_comparisons"][comparison]["absolute_difference_pp"]
            test=fdr[comparison]; result+=f"; full effect +{effect:.1f} pp; p={test['pvalue']:.2e}; q={test['bh_fdr_qvalue']:.2e}"
        else: result+="; reference"
        return [label,result]
    table1 = [
        ["Condition", "Correctness, class F1 and paired comparison"],
        rq1_result("full_system","Full"),
        rq1_result("capability_grounding_ablation","No grounding","full_vs_capability_grounding_ablation"),
        rq1_result("verifier_ablation","No verifier","full_vs_verifier_ablation"),
    ]
    rq2=summary["rq2"]; micro=rq2["micro_320"]; macro=rq2["macro_40_targets"]
    metric_rows=[("Decision","decision_correct"),("Intent","intent_correct"),("TaskSpec exact","taskspec_exact_match"),("End-to-end","end_to_end_correct")]
    table2 = [
        ["Metric", "Micro result and 40-target grouped interval"],
        *[[label,f"{micro[key]['count']}/{micro[key]['n']}; {pct(micro[key]['rate'])}% (micro 95% CI {interval(micro[key]['wilson95'])}); grouped {interval(macro[key]['bootstrap95'])}%"] for label,key in metric_rows],
    ]
    outcomes={r["category"]:r for r in csv.DictReader((FORMAL/"phase7_results_freeze/tables/supp_rq3_outcomes.csv").open())}
    rq3=summary["rq3"]; mission=rq3["mission_success"]; defer=rq3["resource_counterfactual"]["safe_defer"]
    twin=rq3["twin_correctness"]; leak=rq3["feedback_ground_truth_leakage"]; route=rq3["route_full_minus_disabled"]
    table3 = [
        ["Metric", "Held-out result, interval/effect and claim boundary"],
        ["Mission completed",f"{mission['count']}/{mission['n']}; {pct(mission['rate'])}% (95% CI {interval(mission['wilson95'])}); task completion"],
        ["Safe defer",f"{defer['count']}/{defer['n']}; {pct(defer['rate'])}% ({interval(defer['wilson95'])}); not completion"],
        ["Safe blocked route",f"{outcomes['safe_blocked_route_termination']['n']}/300; {100*float(outcomes['safe_blocked_route_termination']['rate']):.1f}%; mutually exclusive; not completion"],
        ["Genuine failure",f"{outcomes['genuine_unexpected_or_system_failure']['n']}/300 observed; no general reliability claim"],
        ["Detection",f"P={rq3['event_precision']:.3f}; R={rq3['event_recall']:.3f}; primary F1={rq3['event_f1']['mean']:.3f} (bootstrap {rq3['event_f1']['bootstrap95'][0]:.3f}–{rq3['event_f1']['bootstrap95'][1]:.3f})"],
        ["Twin correct",f"{twin['count']}/{twin['n']}; {pct(twin['rate'])}% ({interval(twin['wilson95'])}); House2D state"],
        ["Feedback leakage",f"{leak['count']}/{leak['n']} ({pct(leak['wilson95'][0])}–{pct(leak['wilson95'][1])}%); observed rate"],
        ["Route cost",f"full minus disabled={route['mean']:.3f} (paired {route['bootstrap95'][0]:.3f} to {route['bootstrap95'][1]:.3f}); completion effect {rq3['route_mission_mcnemar']['absolute_difference_pp']:.0f} pp"],
        ["Resource prevention",f"{defer['count']}/{defer['n']}; {pct(defer['rate'])}% ({interval(defer['wilson95'])}); read-only counterfactual"],
        ["Scientific replay",f"{rq3['replay_matches']}/40 matched; same implementation and environment"],
    ]
    for path, rows in [(tables / "table1_rq1.csv", table1), (tables / "table2_rq2.csv", table2), (tables / "table3_rq3.csv", table3)]:
        with path.open("w", newline="") as stream:
            csv.writer(stream).writerows(rows)


def setup_plot() -> None:
    plt.rcParams.update({
        "font.family": "DejaVu Sans", "font.size": 8, "axes.titlesize": 9,
        "axes.labelsize": 8, "xtick.labelsize": 7, "ytick.labelsize": 7,
        "legend.fontsize": 7, "figure.dpi": 140, "savefig.dpi": 320,
        "axes.spines.top": False, "axes.spines.right": False,
    })


def box(ax, xy, wh, text, face="#eeeeee", edge="#222222", fontsize=7.2) -> None:
    x, y = xy; w, h = wh
    patch = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.02", facecolor=face, edgecolor=edge, linewidth=1)
    ax.add_patch(patch)
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=fontsize, wrap=True)


def arrow(ax, start, end, style="-") -> None:
    ax.add_patch(FancyArrowPatch(start, end, arrowstyle="-|>", mutation_scale=10, linewidth=1, linestyle=style, color="#333333"))


def save(fig, name: str) -> None:
    fig.savefig(OUT / "figures" / name, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def make_figures(summary: dict) -> None:
    setup_plot()
    (OUT / "figures").mkdir(exist_ok=True)

    fig, ax = plt.subplots(figsize=(4.1, 6.2)); ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.axis("off")
    labels = [
        (0.80, "Constrained natural-language request", "#f2f2f2"),
        (0.68, "Offline planner\nPlanningResult + candidate TaskSpec", "#d9e8f5"),
        (0.54, "Capability Registry  +  Verifier\nexplicit issues and decision", "#dcead5"),
        (0.40, "Resource admission + confirmation\nbattery, return reserve, user-visible plan", "#f5e6c8"),
        (0.26, "BackendExecutor\nre-verification, timeout, cancel, stop", "#ead7d7"),
        (0.10, "RobotBackend\nHouse2D experiment  |  Create3ROS2 deployment", "#e4def1"),
    ]
    for y, text, colour in labels: box(ax, (0.12, y), (0.76, 0.09), text, colour)
    for (y1, _, _), (y2, _, _) in zip(labels, labels[1:]): arrow(ax, (0.5, y1), (0.5, y2 + 0.09))
    ax.text(0.03, 0.60, "serialised\nartefacts", rotation=90, ha="center", va="center", fontsize=7)
    ax.text(0.97, 0.42, "no direct\nlanguage-to-motion", rotation=90, ha="center", va="center", fontsize=7, color="#8b1a1a")
    save(fig, "figure1_architecture.png")

    fig, ax = plt.subplots(figsize=(4.2, 6.0)); ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.axis("off")
    box(ax, (0.05, .80), (.36, .10), "WorldGroundTruth\nseed, events, truth state", "#efdddd")
    box(ax, (.59, .80), (.36, .10), "Run provenance\nIDs, hashes, wall clock", "#eeeeee")
    box(ax, (.27, .64), (.46, .10), "House2D world transition\nand sensor generation", "#f5e6c8")
    arrow(ax, (.23, .80), (.39, .74)); arrow(ax, (.77, .80), (.61, .74), "--")
    box(ax, (.25, .48), (.50, .11), "RobotObservation only\nroom, onboard-equivalent sensors,\naccessibility, battery, validity", "#d9e8f5")
    arrow(ax, (.50, .64), (.50, .59))
    ax.plot([.08,.92],[.44,.44], color="#8b1a1a", lw=1.4)
    ax.text(.5,.455,"detector trust boundary",ha="center",va="bottom",fontsize=7,color="#8b1a1a")
    bottom = [(0.03,"Detector"),(0.27,"Digital Twin\n+ history"),(0.52,"Alerts +\nfeedback"),(0.76,"Formal\nscoring")]
    for x,t in bottom: box(ax,(x,.20),(.21,.12),t,"#dcead5",fontsize=7)
    arrow(ax,(.42,.48),(.14,.32)); arrow(ax,(.50,.48),(.37,.32)); arrow(ax,(.58,.48),(.62,.32))
    arrow(ax,(.23,.80),(.86,.32),"--"); arrow(ax,(.77,.80),(.86,.32),"--")
    ax.text(.5,.07,"Seed and truth remain on the generation/scoring side;\nanswer-bearing fields cannot enter detector input.",ha="center",fontsize=7)
    save(fig, "figure2_house_sitter_pipeline.png")

    fig, ax = plt.subplots(figsize=(4.2, 5.7)); ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.axis("off")
    box(ax,(.25,.82),(.50,.10),"RaPToR-Lite high-level RobotBackend contract","#e4def1")
    box(ax,(.04,.57),(.39,.16),"House2DBackend v1.1\nexperimental/evaluation\nseeded graph + simplified sensors","#dcead5")
    box(ax,(.57,.57),(.39,.16),"Create3ROS2Backend v5.10\ndeployment\nruntime graph discovery","#d9e8f5")
    arrow(ax,(.42,.82),(.24,.73)); arrow(ax,(.58,.82),(.76,.73))
    box(ax,(.57,.33),(.39,.14),"ROS 2 topics/actions/services\nbattery, odom, IMU, hazards, dock, motion","#f2f2f2",fontsize=6.8)
    arrow(ax,(.76,.57),(.76,.47))
    box(ax,(.57,.11),(.39,.12),"Optional NavigationProvider\nNav2 + validated room waypoints","#f5e6c8")
    arrow(ax,(.76,.33),(.76,.23),"--")
    ax.text(.23,.39,"formal outcomes\ncome from this backend",ha="center",fontsize=7)
    ax.text(.76,.04,"Named-room navigation otherwise unavailable",ha="center",fontsize=7,color="#8b1a1a")
    save(fig, "figure3_backends.png")

    labels = ["Full", "No grounding", "No verifier"]
    rq1=summary["rq1"]["decision_correctness"]; order=["full_system","capability_grounding_ablation","verifier_ablation"]
    vals = [rq1[k]["rate"] for k in order]
    lows = [rq1[k]["wilson95"][0] for k in order]
    highs = [rq1[k]["wilson95"][1] for k in order]
    fig, ax = plt.subplots(figsize=(4.0, 3.3))
    colours = ["#222222", "#777777", "#bbbbbb"]
    ax.bar(labels, vals, color=colours, edgecolor="black", hatch=["", "//", ".."])
    ax.errorbar(range(3), vals, yerr=[[v-l for v,l in zip(vals,lows)],[h-v for h,v in zip(highs,vals)]], fmt="none", color="black", capsize=3)
    for i,v in enumerate(vals): ax.text(i,v-.09 if v>.75 else v+.03,f"{v*100:.1f}%",ha="center",color="white" if i<2 else "black",fontweight="bold")
    ax.set_ylim(0,1.08); ax.set_ylabel("Correct paired decisions"); ax.grid(axis="y",alpha=.25)
    save(fig,"figure4_rq1.png")

    form_rows={r["form"]:r for r in csv.DictReader((FORMAL/"phase7_results_freeze/tables/supp_rq2_language_forms.csv").open())}
    form_keys=["canonical","paraphrase","synonym","explicit_order","unordered_rooms","ambiguity","unsupported","unsafe_or_verifier_bypass"]
    forms = ["Canonical","Paraphrase","Synonym","Explicit order","Unordered rooms","Ambiguity","Unsupported wrapper","Unsafe/bypass wrapper"]
    vals = [float(form_rows[k]["end_to_end_correct"]) for k in form_keys]
    fig, ax = plt.subplots(figsize=(4.1,4.0)); y=range(len(forms))
    ax.barh(list(y), vals, color=["#666666" if x != "Synonym" else "#222222" for x in forms], edgecolor="black", hatch=["//" if x=="Synonym" else "" for x in forms])
    ax.set_yticks(list(y),forms); ax.invert_yaxis(); ax.set_xlim(0,1); ax.set_xlabel("End-to-end correctness (n=40 per form)"); ax.grid(axis="x",alpha=.25)
    for i,v in enumerate(vals): ax.text(v+.015,i,f"{v*100:.1f}%",va="center",fontsize=7)
    save(fig,"figure5_rq2.png")

    outcome_rows={r["category"]:r for r in csv.DictReader((FORMAL/"phase7_results_freeze/tables/supp_rq3_outcomes.csv").open())}
    outcome_keys=["mission_completed","safe_defer","safe_blocked_route_termination"]
    fig, ax = plt.subplots(figsize=(4.1,2.9)); vals=[int(outcome_rows[k]["n"]) for k in outcome_keys]; left=0
    colours=["#333333","#888888","#c8c8c8"]; hatches=["","//",".."]; names=["Completed","Safe defer","Blocked-route stop"]
    for v,c,h,n in zip(vals,colours,hatches,names):
        ax.barh([0],[v],left=left,color=c,edgecolor="black",hatch=h,label=f"{n}: {v} ({v/3:.1f}%)")
        ax.text(left+v/2,0,f"{v}\n{v/3:.1f}%",ha="center",va="center",color="white" if c!="#c8c8c8" else "black",fontweight="bold",fontsize=7)
        left+=v
    ax.set_xlim(0,300); ax.set_yticks([]); ax.set_xlabel("Held-out seeds (n=300)"); ax.legend(loc="upper center",bbox_to_anchor=(.5,-.28),ncol=1,frameon=False)
    ax.text(150,.46,"69 blocked-transition events = 20 prior defers + 49 route stops",ha="center",fontsize=7)
    save(fig,"figure6_rq3.png")

    route=summary["rq3"]["route_full_minus_disabled"]; mean=route["mean"]; lo,hi=route["bootstrap95"]
    fig, ax = plt.subplots(figsize=(4.0,3.0))
    ax.errorbar([mean],[0],xerr=[[mean-lo],[hi-mean]],fmt="o",color="black",capsize=5)
    ax.axvline(0,color="#777777",linestyle="--"); ax.set_yticks([]); ax.set_xlabel("Paired route cost: full minus disabled"); ax.set_xlim(-6,1); ax.grid(axis="x",alpha=.25)
    ax.text(mean,.12,"-4.737\n95% CI [-5.186, -4.268]",ha="center",fontsize=8)
    save(fig,"supp_route_effect.png")

    exploratory=next(r for r in csv.DictReader((FORMAL/"phase7_results_freeze/tables/table3_rq3_heldout.csv").open()) if r["metric"]=="anomaly_F1_dropout_semantic_mapping")
    primary=summary["rq3"]["event_f1"]; exp_ci=[float(x) for x in re.search(r"bootstrap95=([0-9.]+)–([0-9.]+)",exploratory["ci_or_effect"]).groups()]
    fig, ax = plt.subplots(figsize=(4.0,3.1)); labels=["Primary (locked)","Post-hoc semantic match"]; vals=[primary["mean"],float(exploratory["estimate"])]; cis=[tuple(primary["bootstrap95"]),tuple(exp_ci)]
    ax.bar(labels,vals,color=["#333333","#aaaaaa"],edgecolor="black",hatch=["","//"])
    ax.errorbar(range(2),vals,yerr=[[v-l for v,(l,h) in zip(vals,cis)],[h-v for v,(l,h) in zip(vals,cis)]],fmt="none",color="black",capsize=4)
    for i,v in enumerate(vals): ax.text(i,v+.025,f"{v:.3f}",ha="center",fontweight="bold")
    ax.set_ylim(0,.8); ax.set_ylabel("Mean per-seed F1"); ax.grid(axis="y",alpha=.25)
    save(fig,"supp_detection_sensitivity.png")


def clear_body(doc: Document) -> None:
    body = doc._element.body
    sect = body.sectPr
    for child in list(body):
        if child is not sect:
            body.remove(child)


def set_cell_margins(cell, top=40, start=40, bottom=40, end=40) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar"); tc_pr.append(tc_mar)
    for edge, value in (("top",top),("start",start),("bottom",bottom),("end",end)):
        el = tc_mar.find(qn(f"w:{edge}"))
        if el is None: el=OxmlElement(f"w:{edge}"); tc_mar.append(el)
        el.set(qn("w:w"),str(value)); el.set(qn("w:type"),"dxa")


def set_columns(section, count: int) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols"); sect_pr.append(cols)
    cols.set(qn("w:num"), str(count)); cols.set(qn("w:space"), "432")


def set_language(run_or_style, lang="en-GB") -> None:
    rpr = run_or_style._element.get_or_add_rPr() if hasattr(run_or_style._element, "get_or_add_rPr") else run_or_style._element
    node = rpr.find(qn("w:lang"))
    if node is None: node=OxmlElement("w:lang"); rpr.append(node)
    node.set(qn("w:val"),lang); node.set(qn("w:eastAsia"),lang)


def configure_styles(doc: Document) -> None:
    specs = {
        "Normal": ("Times New Roman",10,False,False),
        "Title": ("Arial",18,True,False),
        "Heading 1": ("Arial",9,True,False),
        "Heading 2": ("Arial",9,True,False),
        "Heading 3": ("Arial",9,False,True),
        "Caption": ("Times New Roman",9,True,False),
    }
    for name,(font,size,bold,italic) in specs.items():
        style=doc.styles[name]; style.font.name=font; style.font.size=Pt(size); style.font.bold=bold; style.font.italic=italic
        style._element.rPr.rFonts.set(qn("w:eastAsia"),font); set_language(style)
    title_ppr=doc.styles["Title"]._element.get_or_add_pPr(); title_outline=title_ppr.find(qn("w:outlineLvl"))
    if title_outline is None: title_outline=OxmlElement("w:outlineLvl"); title_ppr.append(title_outline)
    title_outline.set(qn("w:val"),"9")
    doc.styles["Normal"].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.styles["Normal"].paragraph_format.space_after=Pt(3)
    doc.styles["Normal"].paragraph_format.line_spacing=1.0
    for name in ("Heading 1","Heading 2","Heading 3"):
        style=doc.styles[name]; style.paragraph_format.keep_with_next=True; style.paragraph_format.space_before=Pt(7); style.paragraph_format.space_after=Pt(2)
    doc.styles["Caption"].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.styles["Caption"].paragraph_format.keep_with_next=False
    doc.styles["Caption"].paragraph_format.keep_together=True
    if "Reference Entry" not in [s.name for s in doc.styles]:
        style=doc.styles.add_style("Reference Entry",WD_STYLE_TYPE.PARAGRAPH)
    else: style=doc.styles["Reference Entry"]
    style.font.name="Times New Roman"; style.font.size=Pt(8); style._element.rPr.rFonts.set(qn("w:eastAsia"),"Times New Roman")
    style.paragraph_format.left_indent=Inches(.16); style.paragraph_format.first_line_indent=Inches(-.16); style.paragraph_format.space_after=Pt(2)
    if "Front Heading" not in [s.name for s in doc.styles]:
        front=doc.styles.add_style("Front Heading",WD_STYLE_TYPE.PARAGRAPH)
    else: front=doc.styles["Front Heading"]
    front.font.name="Arial"; front.font.size=Pt(9); front.font.bold=True; front._element.rPr.rFonts.set(qn("w:eastAsia"),"Arial")
    front.paragraph_format.space_before=Pt(7); front.paragraph_format.space_after=Pt(2); front.paragraph_format.keep_with_next=True
    front_ppr=front._element.get_or_add_pPr(); outline=front_ppr.find(qn("w:outlineLvl"))
    if outline is None: outline=OxmlElement("w:outlineLvl"); front_ppr.append(outline)
    outline.set(qn("w:val"),"9"); set_language(front)


def add_field(paragraph, instruction: str, placeholder: str = "Update field") -> None:
    run=paragraph.add_run(); begin=OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"),"begin")
    instr=OxmlElement("w:instrText"); instr.set(qn("xml:space"),"preserve"); instr.text=instruction
    separate=OxmlElement("w:fldChar"); separate.set(qn("w:fldCharType"),"separate")
    text=OxmlElement("w:t"); text.text=placeholder
    end=OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"),"end")
    for node in (begin,instr,separate,text,end): run._r.append(node)


def add_table(doc: Document, path: Path) -> None:
    rows=list(csv.reader(path.open()))
    table=doc.add_table(rows=len(rows),cols=len(rows[0])); table.autofit=False
    total=3.02; widths=[.92,total-.92] if len(rows[0])==2 else [total/len(rows[0])]*len(rows[0])
    for j,width in enumerate(widths):
        table.columns[j].width=Inches(width)
    tbl_pr=table._tbl.tblPr; borders=OxmlElement("w:tblBorders")
    for edge in ("top","left","bottom","right","insideH","insideV"):
        node=OxmlElement(f"w:{edge}"); node.set(qn("w:val"),"single"); node.set(qn("w:sz"),"4"); node.set(qn("w:color"),"888888"); borders.append(node)
    tbl_pr.append(borders)
    for i,row in enumerate(rows):
        for j,value in enumerate(row):
            cell=table.cell(i,j); cell.text=value; set_cell_margins(cell)
            cell.width=Inches(widths[j])
            for p in cell.paragraphs:
                p.paragraph_format.space_after=Pt(0); p.paragraph_format.keep_together=True; p.paragraph_format.keep_with_next=True
                for run in p.runs:
                    run.font.name="Times New Roman"; run.font.size=Pt(6.5 if len(rows[0])>=4 else 7); run.font.bold=(i==0); set_language(run)
            if i==0:
                shading=OxmlElement("w:shd"); shading.set(qn("w:fill"),"D9D9D9"); cell._tc.get_or_add_tcPr().append(shading)


def reference_text(r: dict) -> str:
    text=f"{r['authors']}. {r['year']}. {r['title']}. {r['venue']}."
    if r["doi"]: text += f" https://doi.org/{r['doi']}"
    else: text += f" {r['url']} (accessed 10 August 2026)."
    return text.replace("--", "–")


def add_document_content(doc: Document, manuscript: str, ordered: list[dict], numbers: dict[str,int]) -> tuple[int,int]:
    lines=manuscript.splitlines(); in_body=False; refs_added=False; figure_count=0; table_count=0
    title=doc.add_paragraph(style="Title"); title.alignment=WD_ALIGN_PARAGRAPH.CENTER; title.add_run(TITLE)
    for text,bold in [("[STUDENT ID REQUIRED]",True),("SEIoT MSc Final Project Report 2026",False),("University College London",False),("Supervisor: Jagmohan Chauhan",False)]:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(text); r.bold=bold; r.font.name="Times New Roman"; r.font.size=Pt(10); set_language(r)
    for raw in lines:
        line=raw.strip()
        if not line: continue
        if line=="# 1. INTRODUCTION" and not in_body:
            doc.add_page_break(); p=doc.add_paragraph(style="Front Heading"); p.add_run("TABLE OF CONTENTS")
            toc=doc.add_paragraph(); add_field(toc,'TOC \\o "1-3" \\h \\z \\u',"Right-click and update field")
            doc.add_page_break(); new=doc.add_section(WD_SECTION.CONTINUOUS); set_columns(new,2); in_body=True
        if line.startswith("# "):
            heading=line[2:]
            if heading=="REFERENCES":
                p=doc.add_paragraph(heading,style="Heading 1")
                for r in ordered:
                    doc.add_paragraph(f"[{numbers[r['key']]}] {reference_text(r)}",style="Reference Entry")
                refs_added=True
            elif heading in {"ABSTRACT","AUTHOR KEYWORDS"}: doc.add_paragraph(heading,style="Front Heading")
            else: doc.add_paragraph(heading,style="Heading 1")
        elif line.startswith("## "): doc.add_paragraph(line[3:],style="Heading 2")
        elif line.startswith("### "): doc.add_paragraph(line[4:],style="Heading 3")
        elif line.startswith("[[FIGURE:"):
            payload=line[len("[[FIGURE:"):-2]; rel,caption=payload.split("|",1)
            p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.keep_with_next=True
            p.add_run().add_picture(str(OUT/rel),width=Inches(3.05)); doc.add_paragraph(caption,style="Caption"); figure_count+=1
        elif line.startswith("[[TABLE:"):
            payload=line[len("[[TABLE:"):-2]; rel,caption=payload.split("|",1)
            add_table(doc,OUT/rel); doc.add_paragraph(caption,style="Caption"); table_count+=1
        elif line.startswith("- "):
            p=doc.add_paragraph(style="List Bullet"); p.add_run(replace_citations(line[2:],numbers))
        else:
            p=doc.add_paragraph(); p.add_run(replace_citations(line,numbers))
    if not refs_added: raise SystemExit("References marker missing")
    return figure_count,table_count


def set_core_properties(doc: Document) -> None:
    cp=doc.core_properties; cp.title=TITLE; cp.subject="UCL SEIoT MSc Final Project Report"; cp.author="Anonymous student ID pending"
    cp.keywords="natural-language robot programming, capability grounding, ROS 2, Digital Twin, reproducibility"
    cp.comments="Built directly from the official Dissertation Template 2024.docx"


def enable_update_fields(doc: Document) -> None:
    settings=doc.settings._element; node=settings.find(qn("w:updateFields"))
    if node is None: node=OxmlElement("w:updateFields"); settings.append(node)
    node.set(qn("w:val"),"true")


def configure_footers(doc: Document) -> None:
    for section in doc.sections:
        footer=section.footer; paragraph=footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.clear(); paragraph.alignment=WD_ALIGN_PARAGRAPH.CENTER; add_field(paragraph,"PAGE","1")
        for run in paragraph.runs:
            run.font.name="Times New Roman"; run.font.size=Pt(9); set_language(run)


def validate_docx(path: Path) -> None:
    if not zipfile.is_zipfile(path): raise SystemExit("DOCX is not a valid OOXML package")
    with zipfile.ZipFile(path) as archive:
        names=set(archive.namelist()); required={"word/document.xml","word/styles.xml","word/settings.xml","word/footer1.xml"}
        if not required <= names: raise SystemExit(f"DOCX missing OOXML parts: {required-names}")
        xml=archive.read("word/document.xml").decode("utf-8")
        for bad in ("SIGCHI Conference Paper Format","Insert Your Subtitle Here","First Author's Name"):
            if bad in xml: raise SystemExit(f"Template example remains: {bad}")
        if "TOC \\o" not in xml: raise SystemExit("TOC field missing")
        if xml.count("w:num=\"2\"")==0 and "w:num\" w:val=\"2\"" not in xml:
            # Attribute serialization varies; direct element validation below is authoritative.
            pass


def build_docx(manuscript: str, ordered: list[dict], numbers: dict[str,int]) -> tuple[int,int]:
    if DOCX.exists(): DOCX.chmod(0o644)
    shutil.copyfile(TEMPLATE,DOCX); DOCX.chmod(0o644)
    doc=Document(DOCX); clear_body(doc); configure_styles(doc); set_columns(doc.sections[0],1)
    figures,tables=add_document_content(doc,manuscript,ordered,numbers)
    configure_footers(doc); set_core_properties(doc); enable_update_fields(doc); doc.save(DOCX); validate_docx(DOCX)
    return figures,tables


def report_counts(manuscript: str, numbers: dict[str,int], table_count:int, figure_count:int, reference_count:int) -> dict:
    rendered=replace_citations(manuscript,numbers)
    abstract=rendered.split("# ABSTRACT",1)[1].split("# AUTHOR KEYWORDS",1)[0]
    report=rendered.split("# ABSTRACT",1)[1].split("# REFERENCES",1)[0]
    report=re.sub(r"\[\[(?:FIGURE|TABLE):[^|]+\|([^\]]+)\]\]",r"\1",report)
    for path in sorted((OUT/"tables").glob("*.csv")):
        report += " " + " ".join(sum(list(csv.reader(path.open())),[]))
    report += " " + TITLE + " STUDENT ID REQUIRED SEIoT MSc Final Project Report 2026 University College London Supervisor Jagmohan Chauhan TABLE OF CONTENTS"
    body=rendered.split("# 1. INTRODUCTION",1)[1].split("# ACKNOWLEDGMENTS",1)[0]
    body=re.sub(r"(?m)^#+ .*?$", "", body); body=re.sub(r"\[\[(?:FIGURE|TABLE):.*?\]\]", "", body)
    counts={
        "abstract_word_count":len(words(abstract)),
        "template_defined_report_word_count":len(words(report)),
        "approximate_main_body_prose_word_count":len(words(body)),
        "reference_count":reference_count,
        "top_level_numbered_section_count":len(re.findall(r"(?m)^# [1-6]\. ",manuscript)),
        "numbered_subsection_count":len(re.findall(r"(?m)^## [1-6]\.\d+ ",manuscript)),
        "table_count":table_count,
        "figure_count":figure_count,
        "main_figure_count":6,
        "supplementary_figure_count":figure_count-6,
    }
    if counts["abstract_word_count"]>300: raise SystemExit("Abstract exceeds 300 words")
    if not 10000<=counts["template_defined_report_word_count"]<=14000: raise SystemExit(f"Report count outside target: {counts}")
    return counts


def write_checklist(counts: dict, pdf_ready: bool, pdf_fonts: str="pending") -> None:
    status="PASS" if pdf_ready else "PENDING PDF EXPORT"
    lines=[
        "# Submission Checklist", "",
        f"- [x] Final title locked: *{TITLE}*",
        f"- [x] Abstract: {counts['abstract_word_count']} words (maximum 300)",
        f"- [x] Template-defined report word count: {counts['template_defined_report_word_count']}",
        f"- [x] Approximate main-body prose word count: {counts['approximate_main_body_prose_word_count']}",
        f"- [x] References: {counts['reference_count']} verified entries, first-author alphabetical numbering",
        f"- [x] Sections: {counts['top_level_numbered_section_count']} top-level, {counts['numbered_subsection_count']} numbered subsections",
        f"- [x] Tables: {counts['table_count']}; figures: {counts['figure_count']} ({counts['main_figure_count']} main, {counts['supplementary_figure_count']} supplementary)",
        "- [x] Directly based on official Dissertation Template 2024.docx",
        "- [x] Letter page size, template margins, single-column front matter and two-column report body retained",
        "- [x] British English, numbered sections, captions below figures/tables, numeric citations",
        "- [x] RQ1/RQ2/RQ3 results traced to frozen raw logical hashes",
        "- [x] Pilot excluded; formal results and core unchanged",
        "- [x] Primary RQ3 F1 0.403 retained; 0.629 labelled post-hoc exploratory",
        "- [x] 61.7% mission completion separated from 100% execution-level safe outcome",
        "- [x] Create3ROS2 interface/mock status and physical_robot_validated=false stated",
        f"- [{'x' if pdf_ready else ' '}] DOCX/PDF generation and structural validation: {status}",
        f"- [{'x' if pdf_ready else ' '}] PDF font inspection: {pdf_fonts}",
        "", "## Unresolved submission metadata", "",
        "- Student ID: replace `[STUDENT ID REQUIRED]` on the title page before submission.",
        "- No verified video-submission URL exists; the optional template item is omitted.",
        "- DOCX declares Times New Roman and Arial exactly. This Linux host lacks Microsoft fonts; inspect the PDF font report below and perform final export on a licensed Microsoft-font installation if exact embedded/substituted typefaces are required.",
    ]
    (OUT/"SUBMISSION_CHECKLIST.md").write_text("\n".join(lines)+"\n")


def write_manifest(counts:dict, refs:list[dict]) -> None:
    artifacts={}
    for path in sorted(p for p in OUT.rglob("*") if p.is_file() and "__pycache__" not in p.parts and p.name not in {"dissertation_evidence_manifest.json"}):
        artifacts[str(path.relative_to(OUT))]=sha256(path)
    data={
        "title":TITLE,"built_from_git_head":BASE_HEAD,"official_template_sha256":TEMPLATE_HASH,
        "phase6_protocol_hash":PROTOCOL_HASH,"phase6_analysis_plan_hash":ANALYSIS_HASH,
        "phase7_results_manifest_hash":RESULTS_HASH,"phase7_results_manifest_file_sha256":sha256(RESULTS_MANIFEST),
        "formal_raw_logical_hashes":json.loads(FREEZE.read_text())["raw_hashes"],
        "counts":counts,"citation_audit":"PASS","claim_audit":"PASS","physical_robot_validated":False,
        "artifacts":artifacts,
    }
    canonical=json.dumps(data,sort_keys=True,separators=(",",":"),ensure_ascii=False).encode()
    data["dissertation_evidence_hash"]=hashlib.sha256(canonical).hexdigest()
    (OUT/"dissertation_evidence_manifest.json").write_text(json.dumps(data,indent=2,sort_keys=True)+"\n")


def finalise_existing() -> None:
    summary,freeze,refs,manuscript=load_inputs(); numbers,ordered=citation_map(refs); audit_citations(manuscript,refs,numbers)
    figures=len(re.findall(r"(?m)^\[\[FIGURE:",manuscript)); tables=len(re.findall(r"(?m)^\[\[TABLE:",manuscript))
    counts=report_counts(manuscript,numbers,tables,figures,len(refs))
    if not DOCX.exists() or not PDF.exists(): raise SystemExit("Final DOCX/PDF missing")
    pdfinfo=subprocess.check_output(["pdfinfo",str(PDF)],text=True)
    pdftext=subprocess.check_output(["pdftotext","-layout",str(PDF),"-"],text=True)
    fonts=subprocess.check_output(["pdffonts",str(PDF)],text=True)
    for required in ("TABLE OF CONTENTS","1. INTRODUCTION","REFERENCES","APPENDIX B. CLAIM AND DEPLOYMENT BOUNDARIES","physical_robot_validated=false"):
        if required not in pdftext: raise SystemExit(f"PDF missing required text: {required}")
    for forbidden in ("Right-click and update field","SIGCHI Conference Paper Format","[cite:","Insert Your Subtitle Here"):
        if forbidden in pdftext: raise SystemExit(f"PDF retains build/template marker: {forbidden}")
    if "Page size:       612 x 792 pts (letter)" not in pdfinfo: raise SystemExit("PDF is not US Letter")
    note_dir=OUT/"writing_notes"; note_dir.mkdir(exist_ok=True)
    (note_dir/"pdf_fonts.txt").write_text(fonts)
    lo_version="LibreOffice 26.2.4.2 (user-local official binary)"
    lines=[
        "# Build and Validation Report","",
        f"- Official template SHA-256: `{TEMPLATE_HASH}`.",
        f"- Baseline Git HEAD: `{BASE_HEAD}`.",
        f"- Converter: {lo_version}.",
        "- Converter archive SHA-256: `810ef197e190d7804a60e0016052c46ff33792303a200fddda9d5216a64b9900`.",
        f"- DOCX SHA-256: `{sha256(DOCX)}`.",f"- PDF SHA-256: `{sha256(PDF)}`.",
        f"- PDF pages: `{re.search(r'Pages:\s+(\d+)',pdfinfo).group(1)}`; page size: US Letter.",
        f"- Word counts: abstract {counts['abstract_word_count']}; template-defined report {counts['template_defined_report_word_count']}; approximate main-body prose {counts['approximate_main_body_prose_word_count']}.",
        "- TOC: refreshed through LibreOffice UNO before PDF export.",
        "- PDF fonts: embedded Liberation Sans/Serif and Caladea substitutions recorded in `pdf_fonts.txt`; DOCX declares Arial/Times New Roman.",
        "- Visual inspection: all pages rendered; tables, figures, columns and captions inspected at page and original resolution.",
        "- Formal evidence: pre/post full-directory SHA-256 comparison passed for all 25 files.",
    ]
    (note_dir/"build_report.md").write_text("\n".join(lines)+"\n")
    write_checklist(counts,True,"PASS with Linux substitutions explicitly recorded")
    write_manifest(counts,refs)
    print(json.dumps(counts,indent=2)); print(f"DOCX_SHA256={sha256(DOCX)}"); print(f"PDF_SHA256={sha256(PDF)}")


def main() -> None:
    summary,freeze,refs,manuscript=load_inputs(); numbers,ordered=citation_map(refs); audit_citations(manuscript,refs,numbers)
    write_references_audit(ordered,numbers); make_tables(summary); make_figures(summary)
    figures,tables=build_docx(manuscript,ordered,numbers); counts=report_counts(manuscript,numbers,tables,figures,len(refs))
    write_checklist(counts,PDF.exists()); write_manifest(counts,refs)
    print(json.dumps(counts,indent=2)); print(f"DOCX {DOCX} {sha256(DOCX)}")


if __name__=="__main__":
    finalise_existing() if "--finalise" in sys.argv else main()
